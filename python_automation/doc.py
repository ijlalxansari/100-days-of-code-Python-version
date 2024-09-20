from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Darul-Karim REC 2024-2025 Timetable', level=1)

# Define the table data
data = [
    ["Grade", "Monday", "Tuesday", "Wednesday", "Thursday", "Saturday", "Sunday", "Teacher"],
    ["Grade 1", "Activity: Craft and Storytelling", "Curriculum + Storytelling", "Quran Recitation + Kalima", "Craft: Kalima Collage", "Dua Recitation + Ginan Singing", "Optional: Play-Based Learning", "[Teacher Name]"],
    ["Grade 2", "Book Reading", "Group Activities + Quiz", "Dua Recitation + Ginan Singing", "Quran Recitation + Kalima", "Interactive Quranic Story", "Presentation + Quiz", "[Teacher Name]"],
    ["Grade 3", "Quran Recitation + Kalima", "Curriculum + Group Discussions", "Dua Recitation + Ginan Meaning", "Quran Recitation + Group Projects", "Optional: Islamic Story Drawing", "Quiz + Discussion", "[Teacher Name]"],
    ["Grade 4", "Activity: Craft + Debate", "Quran Recitation + Kalima", "Curriculum + Debate", "Project: Ismaili History", "Curriculum Review", "Dua & Ginan Interpretation", "[Teacher Name]"],
    ["Grade 5", "Curriculum + Activity", "Group Activity + Quiz", "Dua & Ginan/Farman", "Quran Recitation + Kalima", "Test + Dua Meaning", "Quran Review & Meaning", "[Teacher Name]"],
    ["Grade 6", "Curriculum + Discussion", "Optional: Field Visit", "Curriculum + Quran Recitation", "Kalima Practice + Group Activity", "Ethics Discussion + Group Project", "Dua Recitation + Ginan Meaning", "[Teacher Name]"],
    ["Grade 7", "Quran Recitation + Kalima", "Curriculum + Group Debate", "Dua & Ginan Interpretation", "Class Discussion", "Optional: Leadership Role Play", "Project Presentation", "[Teacher Name]"],
    ["Grade 8", "Curriculum + Activity", "Group Discussion + Quiz", "Quran Recitation + Kalima", "Project Work + Debate", "Interactive Activity + Review", "Reflection + Dua Recitation", "[Teacher Name]"],
    ["Grade 9", "Reflection + Farman Mubarak", "Project Work", "Reflection + Discussion", "Presentation on Ismaili History", "Discussion + Review", "Ginan Practice + Quran Recitation", "[Teacher Name]"],
    ["Grade 10", "Advanced Project Work", "Curriculum Review", "Discussion + Quiz", "Debate + Farman Mubarak", "Group Activity + Presentation", "Class Discussion + Reflection", "[Teacher Name]"],
    ["New Step", "Dua Recitation + Discussion", "Book Reading + Class Discussion", "Optional: Class Activities", "Presentation on Ismaili History", "Discussion + Review", "Ginan Practice + Quran Recitation", "[Teacher Name]"]
]

# Add a table to the document
table = doc.add_table(rows=1, cols=len(data[0]))

# Set the header row
hdr_cells = table.rows[0].cells
for i, header in enumerate(data[0]):
    hdr_cells[i].text = header

# Add the rest of the data
for row_data in data[1:]:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Save the document
doc.save('Darul-Karim_REC_2024-2025_Timetable.docx')

print("Document created successfully.")
